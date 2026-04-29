import streamlit as st
import pandas as pd
import numpy as np
from docx import Document
import os
import zipfile
from pathlib import Path
from copy import deepcopy
from io import BytesIO
from lxml import etree

# =========================================================
# PAGE CONFIG
# =========================================================
st.set_page_config(
    page_title="RDC orders - BOL Generator",
    page_icon="📦",
    layout="wide"
)

# =========================================================
# CONSTANTS / FILE PATHS
# =========================================================
BASE_DIR = Path(__file__).resolve().parent
ASSETS_DIR = BASE_DIR

LTL_QTY_PATH = ASSETS_DIR / "LTL_qty.xlsx"
RDC_LIST_PATH = ASSETS_DIR / "RDC_list.XLSX"
CARRIER_MAP_PATH = ASSETS_DIR / "Carrier List 2025 RDC.xlsx"
CUBE_PATH = ASSETS_DIR / "Cube_mapping.xlsx"
BOL_TEMPLATE_PATH = ASSETS_DIR / "BOL_RDC_template.docx"
MANIFEST_TEMPLATE_PATH = ASSETS_DIR / "Master Manifest_template.docx"

MAX_TEMPLATE_ROWS = 5

# =========================================================
# WORD DOCUMENT PROCESSING FUNCTIONS
# =========================================================
def _replace_in_paragraph(paragraph, replacements):
    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text:
        return

    replaced = full_text
    for k, v in replacements.items():
        replaced = replaced.replace(k, str(v))

    if replaced != full_text and paragraph.runs:
        paragraph.runs[0].text = replaced
        for run in paragraph.runs[1:]:
            run.text = ""


def _replace_in_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _replace_in_paragraph(p, replacements)
            for t in cell.tables:
                _replace_in_table(t, replacements)


def fill_template_doc(template_path, replacements):
    doc = Document(template_path)

    for p in doc.paragraphs:
        _replace_in_paragraph(p, replacements)

    for t in doc.tables:
        _replace_in_table(t, replacements)

    for section in doc.sections:
        for p in section.header.paragraphs:
            _replace_in_paragraph(p, replacements)
        for t in section.header.tables:
            _replace_in_table(t, replacements)

        for p in section.footer.paragraphs:
            _replace_in_paragraph(p, replacements)
        for t in section.footer.tables:
            _replace_in_table(t, replacements)

    return doc


def append_doc_content(master_doc, sub_doc):
    for element in sub_doc.element.body:
        if element.tag.endswith("sectPr"):
            continue
        master_doc.element.body.append(deepcopy(element))


def prepend_doc_content(main_doc, prefix_doc):
    body = main_doc.element.body

    to_insert = [
        deepcopy(el) for el in prefix_doc.element.body
        if not el.tag.endswith("sectPr")
    ]

    pgbrk_xml = (
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:r><w:br w:type="page"/></w:r></w:p>'
    )
    to_insert.append(etree.fromstring(pgbrk_xml))

    for el in reversed(to_insert):
        body.insert(0, el)


# =========================================================
# MANIFEST FUNCTIONS
# =========================================================
def build_po_summary(sid, upload_merged):
    sid_df = upload_merged[upload_merged["SID"] == sid]
    po_summary = (
        sid_df
        .groupby(["SID", "Purchase order no."], as_index=False)
        .agg({
            "Gross weight": "sum",
            "RDC names": "first",
            "Codice RDC": "first",
            "Carrier_mapped": "first",
            "SCAC Code": "first",
            "Cube_po": "sum",
            "Pallet_qty": "sum"
        })
    )
    po_summary["Gross weight"] = po_summary["Gross weight"].astype(int)
    po_summary["Pallet_qty"] = po_summary["Pallet_qty"].astype(int)
    return po_summary


def build_manifest_replacements(sid, po_summary, carrier_name, scac, shipment_rows):
    pickup_date = pd.to_datetime(shipment_rows["Pickup Date"].iloc[0], errors="coerce")
    delivery_date = pickup_date + pd.Timedelta(days=1) if pd.notna(pickup_date) else pd.NaT

    replacements = {
        "{{SID}}": sid,
        "{{CARRIER NAME}}": carrier_name,
        "{{SCAC}}": scac,
        "{{TOTWEIGHT}}": f"{int(shipment_rows['Weight'].iloc[0]):,}",
        "{{TOTCUBE}}": f"{int(shipment_rows['Cube'].iloc[0]):,}",
        "{{PICKUP}}": pickup_date.strftime("%m/%d/%Y") if pd.notna(pickup_date) else "",
        "{{DELIVERY}}": delivery_date.strftime("%m/%d/%Y") if pd.notna(delivery_date) else "",
        "{{RECEIVER}}": "MCDONOUGH IFC",
        **{f"{{{{PO_{i}}}}}": "" for i in range(1, MAX_TEMPLATE_ROWS + 1)},
        **{f"{{{{RDC_{i}}}}}": "" for i in range(1, MAX_TEMPLATE_ROWS + 1)},
        **{f"{{{{Weight_{i}}}}}": "" for i in range(1, MAX_TEMPLATE_ROWS + 1)},
        **{f"{{{{Cube_{i}}}}}": "" for i in range(1, MAX_TEMPLATE_ROWS + 1)},
    }
    return replacements


def _get_row_data_for_po(po_row):
    return {
        "stop": "FINAL",
        "sid": str(po_row["SID"]),
        "po": str(po_row["Purchase order no."]),
        "supplier": "ATLAS CONCORDE - MT PLEASANT",
        "dest": str(po_row["Codice RDC"]),
        "qty": "",
        "haz": "",
        "weight": f"{int(po_row['Gross weight']):,}",
        "cube": f"{int(po_row['Cube_po']):,}",
        "pct_truck": "",
        "case": "",
        "exp_pallets": "",
        "rec_pallets": str(int(po_row["Pallet_qty"]))
    }


def _fill_detail_row(table_row, rd):
    cell_values = [
        rd["stop"], rd["sid"], rd["po"], rd["supplier"], rd["dest"],
        rd["qty"], rd["haz"], rd["weight"], rd["cube"],
        rd["pct_truck"], rd["case"], rd["exp_pallets"], rd["rec_pallets"]
    ]

    cells = table_row.cells

    for j, val in enumerate(cell_values):
        if j < len(cells):
            cell = cells[j]

            if not cell.paragraphs:
                p = cell.add_paragraph()
            else:
                p = cell.paragraphs[0]

            p.text = str(val)

            for extra_p in cell.paragraphs[1:]:
                extra_p._element.getparent().remove(extra_p._element)


def _clear_detail_row(table_row):
    empty = {
        "stop": "", "sid": "", "po": "", "supplier": "", "dest": "",
        "qty": "", "haz": "", "weight": "", "cube": "",
        "pct_truck": "", "case": "", "exp_pallets": "", "rec_pallets": ""
    }
    _fill_detail_row(table_row, empty)


def fill_manifest_table_rows(doc, po_rows_list):
    detail_table = doc.tables[-1]
    data_rows = detail_table.rows[1:]
    n_pos = len(po_rows_list)
    n_template = len(data_rows)

    for i, tr in enumerate(data_rows):
        if i < n_pos:
            _fill_detail_row(tr, po_rows_list[i])
        else:
            _clear_detail_row(tr)

    if n_pos > n_template:
        last_tr = data_rows[-1]
        for i in range(n_template, n_pos):
            new_tr = deepcopy(last_tr._tr)
            detail_table._tbl.append(new_tr)
            new_row = detail_table.rows[-1]
            _fill_detail_row(new_row, po_rows_list[i])


def create_manifest_doc(manifest_template_path, sid, shipment_rows, upload_merged):
    carrier_name = shipment_rows["Carrier_mapped"].iloc[0]
    scac = shipment_rows["SCAC Code"].iloc[0]

    po_summary = build_po_summary(sid, upload_merged)
    replacements = build_manifest_replacements(sid, po_summary, carrier_name, scac, shipment_rows)

    doc = fill_template_doc(manifest_template_path, replacements)

    po_rows_list = [
        _get_row_data_for_po(po_row)
        for _, po_row in po_summary.iterrows()
    ]

    fill_manifest_table_rows(doc, po_rows_list)
    return doc


# =========================================================
# FILE READING HELPERS
# =========================================================
def read_multiple_planex_csv(planex_files):
    cols = {
        "Shipment ID": str,
        "Destination ID": str,
        "Carrier Code": str,
        "Weight": float,
        "Cube": float,
        "Quantity": float
    }

    dfs = []
    for f in planex_files:
        temp = pd.read_csv(f, dtype=cols)
        temp["source_file"] = f.name
        dfs.append(temp)

    df_upload_planex = pd.concat(dfs, ignore_index=True)
    df_upload_planex = df_upload_planex[df_upload_planex["Shipment Status"] == "Accepted"].copy()
    df_upload_planex["Shipment ID"] = df_upload_planex["Shipment ID"].astype(str).str.strip()
    df_upload_planex["Destination ID"] = df_upload_planex["Destination ID"].astype(str).str.strip()
    df_upload_planex["Weight"] = df_upload_planex["Weight"].fillna(0).astype(int)
    df_upload_planex["Cube"] = df_upload_planex["Cube"].fillna(0).astype(int)
    df_upload_planex["Quantity"] = df_upload_planex["Quantity"].fillna(0).astype(int)
    df_upload_planex["Carrier Code"] = df_upload_planex["Carrier Code"].astype(str).str.strip()

    return df_upload_planex


def read_multiple_open_orders_excel(order_files):
    cols_as_str = {
        "Purchase order no.": str,
        "Customer": str,
        "SID": str,
        "DN#": str,
        "Material": str,
        "Carrier": str
    }

    dfs = []
    for f in order_files:
        temp = pd.read_excel(f, dtype=cols_as_str)
        for col in cols_as_str.keys():
            if col in temp.columns:
                temp[col] = temp[col].astype(str).str.strip()
        temp["source_file"] = f.name
        dfs.append(temp)

    df_upload = pd.concat(dfs, ignore_index=True)

    cols_to_drop = [c for c in ["Unnamed: 31", "Unnamed: 32", "Unnamed: 33", "source_file"] if c in df_upload.columns]
    df_upload = df_upload.drop(columns=cols_to_drop, errors="ignore")
    df_upload = df_upload.drop_duplicates()

    df_upload = df_upload[df_upload["SID"].notna()]
    df_upload = df_upload[df_upload["DN#"].notna()]
    df_upload = df_upload.reset_index(drop=True)

    return df_upload


def load_reference_files():
    ltl_qty_df = pd.read_excel(LTL_QTY_PATH)
    ltl_qty_df["SAP Code"] = ltl_qty_df["SAP Code"].astype(str)

    rdc_list_df = pd.read_excel(
        RDC_LIST_PATH,
        dtype={"Zip Code": str, "Cust.(StP)": str, "Codice RDC": str},
        engine="openpyxl"
    )
    for c in ["RDC list", "RDC names", "Adress", "City", "State"]:
        if c in rdc_list_df.columns:
            rdc_list_df[c] = rdc_list_df[c].astype(str).str.strip()

    carrier_df = pd.read_excel(CARRIER_MAP_PATH, dtype={"Code": str}, engine="openpyxl")
    carrier_df.columns = carrier_df.columns.str.strip()
    carrier_df = carrier_df[["SCAC Code", "Carrier", "Code"]]
    carrier_df["SCAC Code"] = carrier_df["SCAC Code"].astype(str).str.strip()
    carrier_df["Carrier"] = carrier_df["Carrier"].astype(str).str.strip()
    carrier_df["Code"] = carrier_df["Code"].astype(str).str.strip()
    carrier_df = carrier_df.dropna(subset=["SCAC Code"])

    df_cube = pd.read_excel(CUBE_PATH, dtype={"Cod_articolo": str}, sheet_name="Mapping")
    df_cube["gross lbs + pallet"] = df_cube["gross lbs + pallet"].round().astype(int)
    df_cube["Cube Vendor PLX"] = df_cube["Cube Vendor PLX"].round().astype(int)
    df_cube = df_cube.rename(columns={"Cube Vendor PLX": "Cube_sku"})

    return ltl_qty_df, rdc_list_df, carrier_df, df_cube


# =========================================================
# CORE PROCESSING
# =========================================================
def process_bol_files(planex_files, order_files):
    ltl_qty_df, rdc_list_df, carrier_df, df_cube = load_reference_files()

    df_upload_planex = read_multiple_planex_csv(planex_files)
    df_upload = read_multiple_open_orders_excel(order_files)

    df_copy = pd.merge(
        df_upload.copy(),
        ltl_qty_df[["SAP Code", "LTL Qty", "Case_Pallet"]],
        left_on="Material",
        right_on="SAP Code",
        how="inner"
    )

    df_copy["Pallet_qty"] = np.ceil(df_copy["Order Quantity"] / df_copy["Case_Pallet"])
    df_copy["Gross weight"] = df_copy["Gross weight"] * 2.20462
    df_copy["Gross weight"] = df_copy["Gross weight"].fillna(0).round().astype(int)
    df_copy["Gross weight"] = df_copy["Gross weight"] + df_copy["Pallet_qty"] * 40

    merged_df = pd.merge(
        df_copy,
        rdc_list_df,
        left_on="Customer",
        right_on="Cust.(StP)",
        how="left"
    )

    df_dest = df_upload_planex[["Shipment ID", "Destination ID", "Pickup Date", "Carrier Code", "Weight", "Cube", "Quantity"]].copy()
    df_dest = df_dest.rename(columns={"Carrier Code": "SCAC Code"})
    df_dest["SCAC Code"] = df_dest["SCAC Code"].astype(str).str.strip()

    # Deduplicate to one row per SID — Weight/Cube/Carrier are shipment-level, not PO-level
    df_dest = df_dest.drop_duplicates(subset=["Shipment ID", "Destination ID"])

    upload_merged = pd.merge(
        merged_df,
        df_dest,
        left_on="SID",
        right_on="Shipment ID",
        how="left"
    )
    
    upload_merged = upload_merged.drop(columns=["Carrier"], errors="ignore")

    upload_merged = pd.merge(
        upload_merged,
        carrier_df[["SCAC Code", "Carrier"]],
        on="SCAC Code",
        how="left"
    )
    upload_merged = upload_merged.rename(columns={"Carrier": "Carrier_mapped"})

    upload_merged = pd.merge(
        upload_merged,
        df_cube[["Cod_articolo", "Cube_sku"]],
        left_on="Material",
        right_on="Cod_articolo",
        how="left"
    )
    upload_merged = upload_merged.drop(columns=["Cod_articolo"], errors="ignore")

    upload_merged["Cube_po"] = upload_merged["Cube_sku"] * upload_merged["Pallet_qty"]
    upload_merged["Cube_po"] = upload_merged["Cube_po"].fillna(0).round().astype(int)

    df_bol = (
        upload_merged
        .groupby(["SID", "RDC names", "Destination ID"], as_index=False)
        .agg({
            "Carrier_mapped": "first",
            "Adress": "first",
            "City": "first",
            "State": "first",
            "Zip Code": "first",
            "SCAC Code": "first",
            "Order Quantity": "sum",
            "Gross weight": "sum",
            "Pallet_qty": "sum",
            "Weight": "first",
            "Cube": "first",
            "Pickup Date": "first",
            "Quantity": "first",
            "Purchase order no.": lambda s: ", ".join(pd.unique(s.astype(str))),
            "DN#": lambda s: ", ".join([x for x in pd.unique(s.astype(str)) if x and x.lower() != "nan"]),
        })
    )

    df_bol["Gross weight"] = df_bol["Gross weight"].astype(int)
    df_bol["Order Quantity"] = df_bol["Order Quantity"].astype(int)
    df_bol["Pallet_qty"] = df_bol["Pallet_qty"].astype(int)

    return df_bol, upload_merged, df_upload, df_upload_planex, df_copy, merged_df


# =========================================================
# DOCX/ZIP OUTPUT
# =========================================================
def generate_documents(df_bol, upload_merged):
    output_buffers = {}

    for shipment_id, shipment_rows in df_bol.groupby("SID"):
        sid = str(shipment_id).strip()
        bol_doc = None

        for _, row in shipment_rows.iterrows():
            replacements = {
                "{{SHIPMENT ID}}": sid,
                "{{PICKUP}}": row.get("Pickup Date", ""),
                "{{CARRIER NAME}}": row.get("Carrier_mapped", ""),
                "{{RDC NAME}}": row.get("RDC names", ""),
                "{{ADRESS}}": row.get("Adress", ""),
                "{{CITY}}": row.get("City", ""),
                "{{STATE}}": row.get("State", ""),
                "{{ZIP CODE}}": row.get("Zip Code", ""),
                "{{SCAC}}": row.get("SCAC Code", ""),
                "{{PO_LIST}}": row.get("Purchase order no.", ""),
                "{{DN_LIST}}": row.get("DN#", ""),
                "{{NUM_PACKAGES}}": row.get("Order Quantity", ""),
                "{{WEIGHT}}": row.get("Gross weight", ""),
                "{{QTY_1}}": row.get("Pallet_qty", ""),
                "{{QTY_PACK}}": row.get("Order Quantity", ""),
            }

            filled_doc = fill_template_doc(BOL_TEMPLATE_PATH, replacements)

            if bol_doc is None:
                bol_doc = filled_doc
            else:
                append_doc_content(bol_doc, filled_doc)

        if shipment_rows["Destination ID"].iloc[0] == "XD16":
            manifest_doc = create_manifest_doc(MANIFEST_TEMPLATE_PATH, sid, shipment_rows, upload_merged)
            prepend_doc_content(bol_doc, manifest_doc)

        bio = BytesIO()
        bol_doc.save(bio)
        bio.seek(0)

        filename = f"Shipment_{sid}.docx"
        output_buffers[filename] = bio.getvalue()

    return output_buffers


def create_zip_from_docs(doc_buffers):
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
        for filename, content in doc_buffers.items():
            zipf.writestr(filename, content)
    zip_buffer.seek(0)
    return zip_buffer


def create_excel_report(df_bol, upload_merged):
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df_bol.to_excel(writer, index=False, sheet_name="BOL Summary")
        upload_merged.to_excel(writer, index=False, sheet_name="Merged Detail")
    output.seek(0)
    return output


def build_missing_bol_reason_table(
    df_upload_planex,
    df_upload,
    df_copy,
    merged_df,
    upload_merged,
    df_bol
):
    accepted_sids = set(df_upload_planex["Shipment ID"].dropna().astype(str).str.strip().unique())
    generated_sids = set(df_bol["SID"].dropna().astype(str).str.strip().unique())
    missing_sids = sorted(accepted_sids - generated_sids)

    open_order_sids = set(df_upload["SID"].dropna().astype(str).str.strip().unique())
    after_ltl_merge_sids = set(df_copy["SID"].dropna().astype(str).str.strip().unique())
    after_planex_merge_sids = set(upload_merged["SID"].dropna().astype(str).str.strip().unique())

    rows = []

    for sid in missing_sids:
        if sid not in open_order_sids:
            reason = "No matching SID in RDC Open Orders"
        elif sid not in after_ltl_merge_sids:
            reason = "Dropped in LTL qty merge (missing material mapping)"
        elif sid not in after_planex_merge_sids:
            reason = "Missing in Planex merge"
        elif sid in after_planex_merge_sids:
            sid_subset = upload_merged[upload_merged["SID"].astype(str).str.strip() == sid]
            if sid_subset["Carrier_mapped"].isna().all():
                reason = "Missing carrier mapping (SCAC not found)"
            else:
                reason = "Present in merged data but no BOL created"
        else:
            reason = "Present in merged data but no BOL created"

        planex_subset = df_upload_planex[
            df_upload_planex["Shipment ID"].astype(str).str.strip() == sid
        ]

        destination_id = ""
        shipment_status = ""
        pickup_date = ""

        if not planex_subset.empty:
            destination_id = planex_subset["Destination ID"].iloc[0] if "Destination ID" in planex_subset.columns else ""
            shipment_status = planex_subset["Shipment Status"].iloc[0] if "Shipment Status" in planex_subset.columns else ""
            pickup_date = planex_subset["Pickup Date"].iloc[0] if "Pickup Date" in planex_subset.columns else ""

        rows.append({
            "Shipment ID": sid,
            "Reason BOL not created": reason,
            "Destination ID": destination_id,
            "Shipment Status": shipment_status,
            "Pickup Date": pickup_date
        })

    return pd.DataFrame(rows)


# =========================================================
# STREAMLIT UI
# =========================================================
st.title("📦 BOL Generator")
st.caption("Generate BOL Word documents and Master Manifests from uploaded files.")

with st.sidebar:
    st.header("Instructions")
    st.markdown("""
    **Upload these files:**
    1. Vendor Planex export CSV file(s)
    2. RDC Open Orders Excel file(s)

    **Bundled in the app:**
    - LTL qty mapping
    - RDC list
    - Carrier mapping
    - Cube mapping
    - Word templates
    """)

st.subheader("Upload input files")

col1, col2 = st.columns(2)

with col1:
    planex_files = st.file_uploader(
        "Vendor Planex CSV files",
        type=["csv"],
        accept_multiple_files=True
    )

with col2:
    order_files = st.file_uploader(
        "RDC Open Orders Excel files",
        type=["xlsx", "xls"],
        accept_multiple_files=True
    )

required_assets = [
    LTL_QTY_PATH,
    RDC_LIST_PATH,
    CARRIER_MAP_PATH,
    CUBE_PATH,
    BOL_TEMPLATE_PATH,
    MANIFEST_TEMPLATE_PATH
]

missing_assets = [p.name for p in required_assets if not p.exists()]

if missing_assets:
    st.error("Missing required files in the app folder:")
    for f in missing_assets:
        st.write(f"- {f}")
    st.stop()

run = st.button("Generate BOL Files", type="primary")

if run:
    if not planex_files:
        st.warning("Please upload at least one Vendor Planex CSV.")
        st.stop()

    if not order_files:
        st.warning("Please upload at least one RDC Open Orders Excel file.")
        st.stop()

    try:
        with st.spinner("Processing files and generating BOLs..."):
            df_bol, upload_merged, df_upload, df_upload_planex, df_copy, merged_df = process_bol_files(planex_files, order_files)
            doc_buffers = generate_documents(df_bol, upload_merged)
            zip_buffer = create_zip_from_docs(doc_buffers)
            excel_report = create_excel_report(df_bol, upload_merged)

        st.success(f"Done. Generated {len(doc_buffers)} shipment document(s).")

        docx_generated = len(doc_buffers)
        unique_sids = df_upload["SID"].dropna().nunique()
        unique_pos = df_upload["Purchase order no."].dropna().nunique()
        unique_dns = df_upload["DN#"].dropna().nunique()

        ifc_pos = (
            upload_merged.loc[upload_merged["Destination ID"] == "XD16", "Purchase order no."]
            .dropna()
            .nunique()
        )

        k1, k2, k3, k4, k5 = st.columns(5)
        k1.metric("DOCX generated", docx_generated)
        k2.metric("Unique SIDs", unique_sids)
        k3.metric("Unique POs", unique_pos)
        k4.metric("Unique Deliveries", unique_dns)
        k5.metric("IFC POs", ifc_pos)

        warning_df = build_missing_bol_reason_table(
            df_upload_planex=df_upload_planex,
            df_upload=df_upload,
            df_copy=df_copy,
            merged_df=merged_df,
            upload_merged=upload_merged,
            df_bol=df_bol
        )

        if not warning_df.empty:
            st.warning(f"{len(warning_df)} accepted Shipment ID(s) did not generate a BOL.")
            st.dataframe(warning_df, use_container_width=True)
        else:
            st.info("All accepted Shipment IDs generated a BOL.")

        ## Warning if Destination ID is different than known ones
        valid_dest = df_bol["Destination ID"].astype(str).str.strip()
        unexpected_dest = df_bol[
            ~valid_dest.eq("XD16") & ~valid_dest.str.match(r"^\d+$")
        ][["SID", "Destination ID", "Carrier_mapped", "RDC names", "Purchase order no."]]

        if not unexpected_dest.empty:
            st.warning(
                f"⚠️ {len(unexpected_dest)} shipment(s) have an unexpected Destination ID "
                f"(neither 'XD16' nor a numeric RDC code):"
            )
            st.dataframe(unexpected_dest, use_container_width=True)

        with st.expander("BOL Summary"):
            st.dataframe(df_bol, use_container_width=True)

        st.download_button(
            label="Download ZIP of BOL files",
            data=zip_buffer,
            file_name="BOL_created.zip",
            mime="application/zip"
        )

        st.download_button(
            label="Download Excel report",
            data=excel_report,
            file_name="BOL_report.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

    except Exception as e:
        st.exception(e)
